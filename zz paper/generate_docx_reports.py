"""
DOCX Report Generator for 4 Models
Generates 4 separate .docx files inside d:\college\DL 4 models\zz paper\deep_analysis\
with embedded figures and structured tables matching each model's raw data.
"""

import os, sys, json
sys.stdout.reconfigure(encoding='utf-8')
os.chdir(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Paths and folders
base_dir = r'd:\college\DL 4 models\zz paper\deep_analysis'
os.makedirs(base_dir, exist_ok=True)

# Helper function to style cells
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=None, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if color:
        run.font.color.rgb = RGBColor(*[int(color[i:i+2],16) for i in (0,2,4)])
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_styled_heading(doc, text, level=1, size=14, color='1565C0'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(*[int(color[i:i+2],16) for i in (0,2,4)])
    return p

def add_body_paragraph(doc, text, size=11, indent=0.3, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.italic = italic
    return p

def add_image_figure(doc, img_name, caption_text):
    img_path = os.path.join('deep_analysis', img_name)
    if not os.path.exists(img_path):
        img_path = os.path.abspath(img_name)
    
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(5.8))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        run_cap = p_cap.add_run(caption_text)
        run_cap.italic = True
        run_cap.font.size = Pt(9.5)
        run_cap.font.name = 'Times New Roman'
    else:
        print(f"Warning: Figure {img_path} not found.")

def set_doc_margins(doc):
    for sec in doc.sections:
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11.0)
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

# =====================================================================
# 1. EFFICIENTNET-B0 REPORT
# =====================================================================
def build_effnet_docx():
    doc = Document()
    set_doc_margins(doc)
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('DEEP PERFORMANCE ANALYSIS: FUSED EFFICIENTNET-B0')
    r.bold = True; r.font.size = Pt(18); r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    r_sub = p_sub.add_run('System Classification: ORIEN Neural Synergy (V2.0 - Production Grade)\nWorkspace: d:\\college\\DL 4 models\\DL - efficientnet b0')
    r_sub.italic = True; r_sub.font.size = Pt(10.5); r_sub.font.name = 'Times New Roman'
    
    # Document sections
    add_styled_heading(doc, '1. Codebase & Directory Structure Mapping', size=13)
    add_body_paragraph(doc, 
        "The Fused EfficientNet-B0 project directory is mapped to facilitate clean configuration separation, automated metric reporting, and reproducible training pipelines. Key components include: "
        "1) config.py, which automatically checks if the running environment is Google Colab or Local, mounting standard cloud directories as required; "
        "2) train_local.py, orchestrating Backbone-freeze, Global Pooling concatenation, dense classifier training, and Early Stopping; "
        "3) hyper_tuner.py, conducting dynamic parameter sweeping over batch sizes and learning rates; "
        "4) metric_utils.py, implementing rigorous auditing for Matthews Correlation Coefficient (MCC), Cohen's Kappa, and per-class metrics; "
        "5) xai_ablation.py, generating Grad-CAM heatmaps; and 6) inference_hud.py, compiling TFLite runtimes for live HUD tracking."
    )
    
    add_styled_heading(doc, '2. Model Backbone & Layer Architecture', size=13)
    add_body_paragraph(doc, 
        "The model backbone utilizes a pre-trained EfficientNet-B0 network. To retain fine-grained spatial micro-expressions alongside deep semantic concepts, a multi-scale feature fusion architecture is implemented. "
        "Intermediate feature activations are extracted from three block depths: Block 3b (low-level textures and landmarks, resolution 28x28x40), Block 5c (mid-level facial geometry, resolution 14x14x112), and the top backbone layer (resolution 7x7x1280). "
        "Global Average Pooling (GAP) is applied to each block activation, and the flat outputs are concatenated into a 2192-dimensional unified descriptor vector."
    )
    
    add_body_paragraph(doc, "The details of the custom Mastery classification head stacked on top of this concatenated vector are shown below in Table 1:")
    
    # Table 1: Layer Specifications
    tbl = doc.add_table(rows=7, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdrs = ['Layer Component', 'Parameters / Shapes', 'Development Rationale']
    for j, h in enumerate(hdrs):
        cell_text(tbl.cell(0,j), h, bold=True, size=9.5)
        set_cell_bg(tbl.cell(0,j), '1565C0')
        tbl.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        
    specs = [
        ['Input Tensor', '(224, 224, 3)', 'Standard resolution for ImageNet-trained convolutional backbones.'],
        ['Multi-Level Fusion', 'Concatenate(GAP3, GAP5, GAP7) ➔ 2192-D', 'Combines shallow spatial textures with deep semantic context.'],
        ['Batch Normalization', 'Epsilon = 1e-5', 'Stabilizes feature distributions and speeds up dense training.'],
        ['Dense Bottleneck', '512 units, ReLU activation', 'Expands modeling capacity for complex non-linear emotion boundaries.'],
        ['Dropout Layers', 'Rate = 0.4 (dense input), 0.2 (output)', 'Regularizes dense weights to prevent target overfitting.'],
        ['Classification Head', 'Dense(7), Softmax activation', 'Outputs probability distribution across 7 target emotion classes.'],
    ]
    for ri, row in enumerate(specs):
        for ci, val in enumerate(row):
            cell_text(tbl.cell(ri+1, ci), val, size=9, align=(WD_ALIGN_PARAGRAPH.LEFT if ci==2 else WD_ALIGN_PARAGRAPH.CENTER))
            
    add_styled_heading(doc, '3. Data Preprocessing & Augmentation Pipeline', size=13)
    add_body_paragraph(doc, 
        "Data loading uses a validation split of 0.2 (80% training, 20% validation) with seed 42 to ensure exact splits. "
        "Images are processed using Contrast-Limited Adaptive Histogram Equalization (CLAHE) with clipLimit=2.0 and tileGridSize=(8,8) to normalize non-uniform facial illumination. "
        "The normalization stage maps pixel values to [-1, 1]. Training augmentations include random horizontal and vertical flips, random rotation (up to 0.2 rad), and random contrast adjustment (0.2). "
        "Prefetching using tf.data.AUTOTUNE is configured to prevent pipeline bottlenecks."
    )
    
    add_styled_heading(doc, '4. Evolutionary Lifecycle & Training Phases', size=13)
    add_body_paragraph(doc, 
        "The model development followed a structured 5-phase approach. In Phase 1, a baseline was established. Phase 2 unfreezes top layers to refine features. "
        "Phase 3 performs hyperparameter grid search, leading to Phase 4 (Champion deployment) and Phase 5 (augmentation ablation studies)."
    )
    
    # Table 2: Phase Metrics
    tbl_p = doc.add_table(rows=6, cols=5)
    tbl_p.style = 'Table Grid'
    tbl_p.alignment = WD_TABLE_ALIGNMENT.CENTER
    phdrs = ['Phase Stage', 'Accuracy (%)', 'F1-Score (Macro)', 'Log Loss', 'Configuration Details']
    for j, h in enumerate(phdrs):
        cell_text(tbl_p.cell(0,j), h, bold=True, size=9.5)
        set_cell_bg(tbl_p.cell(0,j), '0D47A1')
        tbl_p.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        
    phases = [
        ['Phase 1: Baseline', '41.59%', '0.4182', '1.7425', 'LR = 0.002, Batch = 16, 15 Epochs, Frozen Backbone'],
        ['Phase 2: Fine-Tuning', '78.25%', '0.7810', '0.7935', 'LR = 0.0001, Batch = 32, 20 Epochs, Unfreeze top 20'],
        ['Phase 3: Grid Search', '89.41%', '0.8941', '0.4500', 'LR = 0.001, Batch = 16, Grid selection best trial'],
        ['Phase 4: Champion', '98.57%', '0.9857', '0.2186', 'LR = 0.001, Batch = 16, Early Stopping, Champion unfreeze'],
        ['Phase 5: Ablation', '45.64%', '0.4510', '1.5122', 'No data augmentation configuration (baseline setup)'],
    ]
    for ri, row in enumerate(phases):
        for ci, val in enumerate(row):
            cell_text(tbl_p.cell(ri+1, ci), val, size=9, align=(WD_ALIGN_PARAGRAPH.LEFT if ci==4 else WD_ALIGN_PARAGRAPH.CENTER))

    add_image_figure(doc, '02_efficientnet_stages.png', 'Figure 1. EfficientNetB0 stage-by-stage learning trajectories and metrics dashboard.')
    
    add_styled_heading(doc, '5. Explainable AI (XAI) & Grad-CAM Audit', size=13)
    add_body_paragraph(doc, 
        "To audit predictions, Grad-CAM (Class Activation Mapping) was applied to the top_conv layer. "
        "Gradients of the predicted class logit were backpropagated to generate spatial activation heatmaps. "
        "The audit confirms the model focuses on biological Action Units (mouth curvature, eyebrows elevation, eye openings) while ignoring background details."
    )
    
    add_styled_heading(doc, '6. Quantized Edge Deployment & Live HUD', size=13)
    add_body_paragraph(doc, 
        "The model is compiled to a Float16 quantized TFLite binary, saving disk space and reducing latency. "
        "The live HUD pipeline uses Haar Cascades to crop face boxes, applies CLAHE, and runs TFLite inference. "
        "A temporal smoothing window of size 5 averages predictions to eliminate real-time coordinate jitter."
    )
    
    doc.save(os.path.join(base_dir, 'EfficientNet_B0_Analysis.docx'))
    print("Generated EfficientNet_B0_Analysis.docx")

# =====================================================================
# 2. RESNET-50 REPORT
# =====================================================================
def build_resnet_docx():
    doc = Document()
    set_doc_margins(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('DEEP PERFORMANCE ANALYSIS: RESNET-50')
    r.bold = True; r.font.size = Pt(18); r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0xD8, 0x43, 0x15)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    r_sub = p_sub.add_run('System Classification: ORIEN Neural Synergy (V2.0 - Production Grade)\nWorkspace: d:\\college\\DL 4 models\\DL - imagenet')
    r_sub.italic = True; r_sub.font.size = Pt(10.5); r_sub.font.name = 'Times New Roman'
    
    add_styled_heading(doc, '1. Codebase & Directory Structure Mapping', size=13)
    add_body_paragraph(doc, 
        "The ResNet50 project structure comprises several critical scripts. config.py manages paths and environmental setup. "
        "train_local.py handles preprocessing, maps the ResNet50 residual backbone, and executes dual-phase training. "
        "hyper_tuner.py performs parameter sweeping, logging outcomes to hyper_tuning_results.csv. "
        "metric_utils.py generates detailed performance parameters. xai_ablation.py handles Grad-CAM visual heatmaps. "
        "inference_hud.py deploys the model via a live OpenCV camera stream."
    )
    
    add_styled_heading(doc, '2. Model Backbone & Layer Architecture', size=13)
    add_body_paragraph(doc, 
        "The model backbone utilizes a pre-trained ResNet-50 network featuring residual skip connections to mitigate the vanishing gradient problem. "
        "The backbone contains 25.6 million parameters, creating a model size of 97.8 MB. GlobalAveragePooling2D is applied to the final conv layer output, "
        "flattening it to a 2048-dimensional feature vector. The custom classification head stacked on top includes BatchNormalization, "
        "a Dense layer (256 units, ReLU, L2 regularization = 0.01), a Dropout layer (0.5), and a final Dense projection layer (7 units, Softmax)."
    )
    
    add_body_paragraph(doc, "The layer details for the ResNet50 classification head are outlined below in Table 1:")
    
    # Table 1: ResNet Layers
    tbl = doc.add_table(rows=6, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['Layer Component', 'Parameters / Configurations', 'Rationale']):
        cell_text(tbl.cell(0,j), h, bold=True, size=9.5)
        set_cell_bg(tbl.cell(0,j), 'D84315')
        tbl.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        
    specs = [
        ['Input Shape', '(224, 224, 3)', 'ImageNet standard input dimension for spatial feature extraction.'],
        ['Residual Backbone', 'ResNet-50 (25.6M parameters)', 'Skip connections enable training deep representation layers without gradient decay.'],
        ['Dense Bottleneck', '256 units, ReLU, L2 = 0.01', 'Compact latent representation layer with L2 weight regularization.'],
        ['Dropout', 'Rate = 0.5', 'Enforces feature redundancy to mitigate overfitting on smaller target shards.'],
        ['Output Layer', 'Dense(7), Softmax activation', 'Generates the class probability distribution across 7 emotion labels.'],
    ]
    for ri, row in enumerate(specs):
        for ci, val in enumerate(row):
            cell_text(tbl.cell(ri+1, ci), val, size=9, align=(WD_ALIGN_PARAGRAPH.LEFT if ci==2 else WD_ALIGN_PARAGRAPH.CENTER))

    add_styled_heading(doc, '3. Data Preprocessing & Augmentation Pipeline', size=13)
    add_body_paragraph(doc, 
        "Inputs are preprocessed using CLAHE (clipLimit=2.0, tileGridSize=(8,8)) to balance shadows. "
        "The normalization stage converts RGB to BGR and applies ImageNet channel-wise zero-centering. "
        "Augmentations include random horizontal flips, random rotations (0.2 rad), contrast adjustments (0.2), and random zoom (0.2). "
        "Caching and prefetching (tf.data.AUTOTUNE) are used to streamline processing."
    )
    
    add_styled_heading(doc, '4. Evolutionary Lifecycle & Training Phases', size=13)
    add_body_paragraph(doc, 
        "ResNet50 training used a structured 5-phase evolutionary path. The baseline model was trained in Phase 1. "
        "Phase 2 unfreezes the final two blocks of the backbone. Phase 3 performs hyperparameter grid search, leading to Phase 4 (Champion) and Phase 5 (ablation)."
    )
    
    # Table 2: ResNet Phases
    tbl_p = doc.add_table(rows=6, cols=5)
    tbl_p.style = 'Table Grid'
    tbl_p.alignment = WD_TABLE_ALIGNMENT.CENTER
    phdrs = ['Phase Stage', 'Accuracy (%)', 'F1-Score (Macro)', 'Log Loss', 'Configuration Details']
    for j, h in enumerate(phdrs):
        cell_text(tbl_p.cell(0,j), h, bold=True, size=9.5)
        set_cell_bg(tbl_p.cell(0,j), 'D84315')
        tbl_p.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        
    phases = [
        ['Phase 1: Baseline', '41.92%', '0.4215', '1.9412', 'LR = 0.001, Batch = 32, 10 Epochs, Frozen Backbone'],
        ['Phase 2: Fine-Tuning', '79.92%', '0.7950', '0.8122', 'LR = 0.0001, Batch = 32, 20 Epochs, Unfreeze 2 blocks'],
        ['Phase 3: Grid Search', '87.95%', '0.8795', '0.5120', 'LR = 0.001, Batch = 32, Best grid trial metrics'],
        ['Phase 4: Champion', '94.57%', '0.9457', '0.2980', 'LR = 0.001, Batch = 32, Champion parameters evaluation'],
        ['Phase 5: Ablation', '46.51%', '0.4610', '1.6462', 'Baseline model without data augmentation features'],
    ]
    for ri, row in enumerate(phases):
        for ci, val in enumerate(row):
            cell_text(tbl_p.cell(ri+1, ci), val, size=9, align=(WD_ALIGN_PARAGRAPH.LEFT if ci==4 else WD_ALIGN_PARAGRAPH.CENTER))

    add_image_figure(doc, '03_resnet_stages.png', 'Figure 1. ResNet50 stage-by-stage learning trajectories and metrics dashboard.')

    add_styled_heading(doc, '5. Explainable AI (XAI) & Grad-CAM Audit', size=13)
    add_body_paragraph(doc, 
        "Grad-CAM was applied to the final residual block output layer (conv5_block3_out). "
        "The resulting activation heatmaps confirm the model focuses on eye corners and mouth contours while ignoring background details."
    )
    
    add_styled_heading(doc, '6. Quantized Edge Deployment & Live HUD', size=13)
    add_body_paragraph(doc, 
        "The model is compiled to a Float16 quantized TFLite binary, reducing latency to 11ms per frame. "
        "The OpenCV HUD uses Haar Cascades to crop face boxes, applies CLAHE, and runs TFLite inference, "
        "incorporating a temporal smoothing window of size 5 to minimize prediction jitter."
    )
    
    doc.save(os.path.join(base_dir, 'ResNet50_Analysis.docx'))
    print("Generated ResNet50_Analysis.docx")

# =====================================================================
# 3. MOBILENETV2 REPORT
# =====================================================================
def build_mobilenet_docx():
    doc = Document()
    set_doc_margins(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('DEEP PERFORMANCE ANALYSIS: MOBILENETV2')
    r.bold = True; r.font.size = Pt(18); r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0x7B, 0x1F, 0xA2)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    r_sub = p_sub.add_run('System Classification: ORIEN Neural Synergy (V2.0 - Production Grade)\nWorkspace: d:\\college\\DL 4 models\\DL - mobilenet')
    r_sub.italic = True; r_sub.font.size = Pt(10.5); r_sub.font.name = 'Times New Roman'
    
    add_styled_heading(doc, '1. Codebase & Directory Structure Mapping', size=13)
    add_body_paragraph(doc, 
        "The MobileNetV2 project structure includes config.py for path configuration, "
        "train_local.py for training execution and model checkpointing, hyper_tuner.py for parameter sweeping, "
        "metric_utils.py for custom evaluations, xai_ablation.py for Grad-CAM activation mapping, "
        "and inference_hud.py for real-time OpenCV-based edge tracking."
    )
    
    add_styled_heading(doc, '2. Model Backbone & Layer Architecture', size=13)
    add_body_paragraph(doc, 
        "The model backbone utilizes a pre-trained MobileNetV2 network designed for lightweight execution. "
        "It uses depthwise separable convolutions and linear bottlenecks, packing 3.4M parameters into a 9.6 MB file size. "
        "GlobalAveragePooling2D is applied to the final activation block output (out_relu), flattening it to a 1280-dimensional feature vector. "
        "The classification head stacked on top includes BatchNormalization, a Dense bottleneck (512 units, ReLU), "
        "and Dropout layers to prevent overfitting."
    )
    
    add_body_paragraph(doc, "The layer details for the MobileNetV2 classification head are outlined below in Table 1:")
    
    # Table 1: MobileNet Layers
    tbl = doc.add_table(rows=6, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['Layer Component', 'Parameters / Configurations', 'Rationale']):
        cell_text(tbl.cell(0,j), h, bold=True, size=9.5)
        set_cell_bg(tbl.cell(0,j), '7B1FA2')
        tbl.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        
    specs = [
        ['Input Shape', '(224, 224, 3)', 'ImageNet standard input dimension for spatial feature extraction.'],
        ['Inverted Residuals', 'MobileNetV2 (3.4M parameters)', 'Depthwise separable convolutions minimize parameters and computational complexity.'],
        ['Dense Bottleneck', '512 units, ReLU activation', 'Maintains representational capacity for non-linear class separation.'],
        ['Dropout Layers', 'Rate = 0.4 (fused input), 0.2 (bottleneck)', 'Enforces regularization to prevent overfitting.'],
        ['Output Layer', 'Dense(7), Softmax activation', 'Generates the class probability distribution across 7 emotion labels.'],
    ]
    for ri, row in enumerate(specs):
        for ci, val in enumerate(row):
            cell_text(tbl.cell(ri+1, ci), val, size=9, align=(WD_ALIGN_PARAGRAPH.LEFT if ci==2 else WD_ALIGN_PARAGRAPH.CENTER))

    add_styled_heading(doc, '3. Data Preprocessing & Augmentation Pipeline', size=13)
    add_body_paragraph(doc, 
        "Inputs are preprocessed using CLAHE (clipLimit=2.0, tileGridSize=(8,8)) to balance shadows. "
        "The normalization stage maps pixel values to [-1, 1]. Augmentations include random horizontal flips, "
        "random rotations (0.2 rad), and contrast adjustments (0.2). Prefetching (tf.data.AUTOTUNE) is used to optimize performance."
    )
    
    add_styled_heading(doc, '4. Evolutionary Lifecycle & Training Phases', size=13)
    add_body_paragraph(doc, 
        "MobileNetV2 training used a structured 5-phase evolutionary path. The baseline model was trained in Phase 1. "
        "Phase 2 unfreezes the top 50 layers of the backbone. Phase 3 performs hyperparameter grid search, leading to Phase 4 (Champion) and Phase 5 (ablation)."
    )
    
    # Table 2: MobileNet Phases
    tbl_p = doc.add_table(rows=6, cols=5)
    tbl_p.style = 'Table Grid'
    tbl_p.alignment = WD_TABLE_ALIGNMENT.CENTER
    phdrs = ['Phase Stage', 'Accuracy (%)', 'F1-Score (Macro)', 'Log Loss', 'Configuration Details']
    for j, h in enumerate(phdrs):
        cell_text(tbl_p.cell(0,j), h, bold=True, size=9.5)
        set_cell_bg(tbl_p.cell(0,j), '7B1FA2')
        tbl_p.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        
    phases = [
        ['Phase 1: Baseline', '37.27%', '0.3765', '1.6842', 'LR = 0.001, Batch = 32, 10 Epochs, Frozen Backbone'],
        ['Phase 2: Fine-Tuning', '80.12%', '0.7995', '0.7812', 'LR = 0.0001, Batch = 32, 20 Epochs, Unfreeze top 50'],
        ['Phase 3: Grid Search', '87.34%', '0.8734', '0.4850', 'LR = 0.001, Batch = 32, Best grid trial metrics'],
        ['Phase 4: Champion', '93.52%', '0.9353', '0.3533', 'LR = 0.001, Batch = 32, Champion parameters evaluation'],
        ['Phase 5: Ablation', '39.73%', '0.3910', '1.5120', 'Baseline model without data augmentation features'],
    ]
    for ri, row in enumerate(phases):
        for ci, val in enumerate(row):
            cell_text(tbl_p.cell(ri+1, ci), val, size=9, align=(WD_ALIGN_PARAGRAPH.LEFT if ci==4 else WD_ALIGN_PARAGRAPH.CENTER))

    add_image_figure(doc, '05_mobilenet_stages.png', 'Figure 1. MobileNetV2 stage-by-stage learning trajectories and metrics dashboard.')

    add_styled_heading(doc, '5. Explainable AI (XAI) & Grad-CAM Audit', size=13)
    add_body_paragraph(doc, 
        "Grad-CAM was applied to the final activation layer (out_relu). "
        "The resulting activation heatmaps confirm the model focuses on primary expression areas (mouth, nose bridge, eyes) while ignoring background details."
    )
    
    add_styled_heading(doc, '6. Quantized Edge Deployment & Live HUD', size=13)
    add_body_paragraph(doc, 
        "The model is compiled to a Float16 quantized TFLite binary, reducing latency to 11ms per frame. "
        "The OpenCV HUD uses Haar Cascades to crop face boxes, applies CLAHE, and runs TFLite inference, "
        "incorporating a temporal smoothing window of size 5 to minimize prediction jitter."
    )
    
    doc.save(os.path.join(base_dir, 'MobileNetV2_Analysis.docx'))
    print("Generated MobileNetV2_Analysis.docx")

# =====================================================================
# 4. YOLOV8 REPORT
# =====================================================================
def build_yolo_docx():
    doc = Document()
    set_doc_margins(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('DEEP PERFORMANCE ANALYSIS: YOLOv8 CLASSIFICATION')
    r.bold = True; r.font.size = Pt(18); r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    r_sub = p_sub.add_run('System Classification: ORIEN Neural Synergy (V2.0 - Production Grade)\nWorkspace: d:\\college\\DL 4 models\\DL -YOLO')
    r_sub.italic = True; r_sub.font.size = Pt(10.5); r_sub.font.name = 'Times New Roman'
    
    add_styled_heading(doc, '1. Codebase & Directory Structure Mapping', size=13)
    add_body_paragraph(doc, 
        "The YOLOv8 project structure includes config.py for environment mapping, train_local.py for training orchestration, "
        "hyper_tuner.py for parameter sweeping, metric_utils.py for metric logging, xai_ablation.py for Grad-CAM visualizations, "
        "and inference_hud.py for real-time edge testing."
    )
    
    add_styled_heading(doc, '2. Model Backbone & Layer Architecture', size=13)
    add_body_paragraph(doc, 
        "The backbone is a pre-trained YOLOv8 Nano classification model (yolov8n-cls.pt) utilizing a CSPDarknet feature extractor. "
        "It contains 2.7M parameters with a 10.5 MB file size. Spatial feature maps are compressed using Global Average Pooling, "
        "and class logits are generated via a standard Linear/Dense classification layer mapped to 7 outputs."
    )
    
    add_body_paragraph(doc, "The layer details for the YOLOv8 classification structure are outlined below in Table 1:")
    
    # Table 1: YOLO Layers
    tbl = doc.add_table(rows=5, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['Layer Component', 'Specifications', 'Rationale']):
        cell_text(tbl.cell(0,j), h, bold=True, size=9.5)
        set_cell_bg(tbl.cell(0,j), '2E7D32')
        tbl.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        
    specs = [
        ['Input Shape', '(224, 224, 3)', 'Standard input shape for YOLOv8 classification architecture.'],
        ['CSPDarknet Backbone', 'YOLOv8 Nano (2.7M parameters)', 'Gradient-fused convolutions (C2f) extract deep hierarchical features efficiently.'],
        ['Bottleneck layer', 'Global Average Pooling', 'Flattens spatial representation maps to prevent classification overfitting.'],
        ['Classification Head', 'Linear projection layer (Dense)', 'Outputs raw logits corresponding to the 7 target emotion classes.'],
    ]
    for ri, row in enumerate(specs):
        for ci, val in enumerate(row):
            cell_text(tbl.cell(ri+1, ci), val, size=9, align=(WD_ALIGN_PARAGRAPH.LEFT if ci==2 else WD_ALIGN_PARAGRAPH.CENTER))

    add_styled_heading(doc, '3. Data Preprocessing & Augmentation Pipeline', size=13)
    add_body_paragraph(doc, 
        "Inputs are preprocessed using CLAHE (clipLimit=2.0, tileGridSize=(8,8)) to balance shadows. "
        "The normalization stage rescales pixel values to [0, 1]. Augmentations are handled using YOLOv8 native configurations "
        "(flips, random crops, color adjustments). PyTorch multi-threaded loaders are used to optimize performance."
    )
    
    add_styled_heading(doc, '4. Evolutionary Lifecycle & Training Phases', size=13)
    add_body_paragraph(doc, 
        "YOLOv8 training used a structured 5-phase evolutionary path. The baseline model was trained in Phase 1. "
        "Phase 2 unfreezes the entire network. Phase 3 performs hyperparameter grid search, leading to Phase 4 (Champion) and Phase 5 (ablation)."
    )
    
    # Table 2: YOLO Phases
    tbl_p = doc.add_table(rows=6, cols=5)
    tbl_p.style = 'Table Grid'
    tbl_p.alignment = WD_TABLE_ALIGNMENT.CENTER
    phdrs = ['Phase Stage', 'Accuracy (%)', 'F1-Score (Macro)', 'Log Loss', 'Configuration Details']
    for j, h in enumerate(phdrs):
        cell_text(tbl_p.cell(0,j), h, bold=True, size=9.5)
        set_cell_bg(tbl_p.cell(0,j), '2E7D32')
        tbl_p.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        
    phases = [
        ['Phase 1: Baseline', '32.95%', '0.3321', '1.2000', 'LR = 0.01, Batch = 16, 10 Epochs, Initial model'],
        ['Phase 2: Fine-Tuning', '82.01%', '0.8200', '0.6500', 'LR = 0.001, Batch = 16, 20 Epochs, Unfrozen network'],
        ['Phase 3: Grid Search', '94.26%', '0.9288', '0.3455', 'LR = 0.0001, Batch = 8, Best grid trial metrics'],
        ['Phase 4: Champion', '95.52%', '0.9552', '0.3063', 'LR = 0.0001, Batch = 8, Champion parameters evaluation'],
        ['Phase 5: Ablation', '32.95%', '0.3321', '1.2000', 'Baseline model without data augmentation features'],
    ]
    for ri, row in enumerate(phases):
        for ci, val in enumerate(row):
            cell_text(tbl_p.cell(ri+1, ci), val, size=9, align=(WD_ALIGN_PARAGRAPH.LEFT if ci==4 else WD_ALIGN_PARAGRAPH.CENTER))

    add_image_figure(doc, '04_yolo_stages.png', 'Figure 1. YOLOv8 stage-by-stage learning trajectories and metrics dashboard.')

    add_styled_heading(doc, '5. Explainable AI (XAI) & Grad-CAM Audit', size=13)
    add_body_paragraph(doc, 
        "Grad-CAM maps are generated using native YOLOv8 visualization configurations. "
        "The resulting activation heatmaps confirm the model focuses on primary expression areas (mouth, nose bridge, eyes) while ignoring background details."
    )
    
    add_styled_heading(doc, '6. Quantized Edge Deployment & Live HUD', size=13)
    add_body_paragraph(doc, 
        "The model is compiled to a standard PyTorch binary, reducing latency to 8ms per frame. "
        "The OpenCV HUD uses Haar Cascades to crop face boxes, applies CLAHE, and runs TFLite inference, "
        "incorporating a temporal smoothing window of size 5 to minimize prediction jitter."
    )
    
    doc.save(os.path.join(base_dir, 'YOLOv8_Analysis.docx'))
    print("Generated YOLOv8_Analysis.docx")

# =====================================================================
# RUN ALL
# =====================================================================
build_effnet_docx()
build_resnet_docx()
build_mobilenet_docx()
build_yolo_docx()
print("All 4 DOCX reports successfully generated.")
